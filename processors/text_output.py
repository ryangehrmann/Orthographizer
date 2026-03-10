"""
processors/text_output.py
Generate output files from a processed list of paragraphs:

  1. IPA-only (.docx)    — Khmer tokens replaced with IPA (and Lao when available);
                           ambiguous wrapped in {}
  2. Interlinear (.html) — each Khmer word shown with its IPA gloss below,
                           and Lao gloss when the pipeline produces it;
                           using CSS flexbox stacking; ambiguous IPA in red + {}
"""
import html as _html
import io
from docx import Document
from docx.shared import RGBColor, Cm
from core.modifiers import syllables_to_word
from core.lao_assembly import syllables_to_lao


def _make_document() -> Document:
    doc = Document()
    # Remove default margins for a compact layout
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return doc


def _transcribe_token(
    token: str,
    is_khmer: bool,
    pipeline_run_fn,
    flat_dict: dict,
    segments_dict: dict,
) -> tuple[str, bool, str]:
    """
    Transcode a token if it contains Khmer text.

    Returns:
        (ipa_text, is_ambiguous, lao_text) — ipa_text is the transcoded IPA
        string (or the original token for non-Khmer); is_ambiguous is True if
        any syllable was flagged ambiguous; lao_text is the assembled Lao-script
        form (empty string when the pipeline does not produce Lao output).
    """
    if not is_khmer:
        return token, False, ""

    ipa, syllables = pipeline_run_fn(token, flat_dict, segments_dict)
    if ipa is None:
        return token, False, ""

    is_ambiguous = any(s.get("ambiguous", False) for s in syllables)
    # Reconstruct the clean word form from modifier-processed syllable dicts
    # rather than using the raw IPA string (which still contains S/T/W/¹/²).
    word = syllables_to_word(syllables) if syllables else ipa
    lao = syllables_to_lao(syllables) if syllables else ""
    return word, is_ambiguous, lao


def _format_ipa(ipa: str, ambiguous: bool) -> str:
    if ambiguous:
        return "{" + ipa + "}"
    return ipa


# ── IPA-only output ──────────────────────────────────────────────────────────

def build_ipa_only_docx(
    paragraphs: list[list[tuple[str, bool]]],
    pipeline_run_fn,
    flat_dict: dict,
    segments_dict: dict,
) -> bytes:
    """
    Build a .docx where each paragraph of Khmer text is replaced with IPA.
    Ambiguous sequences are wrapped in {}.
    """
    doc = _make_document()

    for para_tokens in paragraphs:
        if not para_tokens:
            doc.add_paragraph()
            continue

        para = doc.add_paragraph()
        for token, is_khmer in para_tokens:
            if token.strip() == "":
                # Preserve whitespace as plain text run
                para.add_run(token)
            else:
                ipa, ambiguous, lao = _transcribe_token(
                    token, is_khmer, pipeline_run_fn, flat_dict, segments_dict
                )
                if not is_khmer:
                    para.add_run(token)
                else:
                    text = _format_ipa(ipa, ambiguous)
                    if lao:
                        text = f"{text} ({lao})"
                    run = para.add_run(text)
                    if ambiguous:
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Interlinear HTML output ──────────────────────────────────────────────────

_CSS = """
body {
    font-family: 'Segoe UI', 'Noto Sans Khmer', Arial, sans-serif;
    font-size: 19px;
    padding: 2em 3em;
    background: #fff;
    color: #222;
    max-width: 960px;
    margin: 0 auto;
}
p {
    margin: 0 0 1.8em 0;
}
.word {
    display: inline-flex;
    flex-direction: column;
    align-items: center;
    margin: 0 0.25em 0.25em 0;
    vertical-align: top;
}
.orig {
    font-size: 1em;
    line-height: 1.6;
}
.gloss {
    font-size: 0.82em;
    color: #555;
    line-height: 1.5;
    white-space: nowrap;
}
.lao {
    font-family: 'Saysettha OT', 'Phetsarath OT', 'Noto Serif Lao', serif;
    font-size: 1.02em;
    color: #2255aa;
    line-height: 1.5;
    white-space: nowrap;
}
.ambig .gloss {
    color: #cc0000;
    font-weight: bold;
}
"""

_GOOGLE_FONTS_LINK = (
    '<link rel="preconnect" href="https://fonts.googleapis.com">'
    '<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>'
    '<link href="https://fonts.googleapis.com/css2?family=Noto+Serif+Lao&display=swap" rel="stylesheet">'
)


def build_interlinear_html(
    paragraphs: list[list[tuple[str, bool]]],
    pipeline_run_fn,
    flat_dict: dict,
    segments_dict: dict,
) -> bytes:
    """
    Build a .html file with word-spaced interlinear layout using CSS flexbox.
    Each Khmer word is stacked above its IPA gloss. Non-Khmer words appear
    inline with an empty gloss row so vertical alignment stays consistent
    within mixed paragraphs. Purely non-Khmer paragraphs are plain text.
    Ambiguous IPA is shown in red and wrapped in {}.
    """
    parts = [
        "<!DOCTYPE html>",
        '<html lang="km">',
        "<head>",
        '  <meta charset="UTF-8">',
        "  <title>Interlinear \u2013 Orthographizer</title>",
        f"  {_GOOGLE_FONTS_LINK}",
        f"  <style>{_CSS}  </style>",
        "</head>",
        "<body>",
    ]

    for para_tokens in paragraphs:
        word_tokens = [(t, k) for t, k in para_tokens if t.strip()]

        if not word_tokens:
            parts.append("<p>&nbsp;</p>")
            continue

        has_khmer = any(k for _, k in word_tokens)

        if not has_khmer:
            # Plain paragraph — no interlinear needed
            text = _html.escape(" ".join(t for t, _ in word_tokens))
            parts.append(f"<p>{text}</p>")
            continue

        # Mixed or all-Khmer paragraph — build stacked word spans
        # Detect whether any token in this paragraph has Lao output
        # (used to decide whether to reserve a Lao row for all tokens)
        has_lao = False
        transcribed = []
        for token, is_khmer in word_tokens:
            if is_khmer:
                ipa, ambiguous, lao = _transcribe_token(
                    token, is_khmer, pipeline_run_fn, flat_dict, segments_dict
                )
                if lao:
                    has_lao = True
            else:
                ipa, ambiguous, lao = token, False, ""
            transcribed.append((token, is_khmer, ipa, ambiguous, lao))

        spans = []
        for token, is_khmer, ipa, ambiguous, lao in transcribed:
            orig_esc = _html.escape(token)
            if not is_khmer:
                # Non-Khmer token: text always at top, placeholder rows below.
                # (Keeps numbers/punctuation from sinking to the bottom of a
                # 3-row stack when has_lao=True.)
                if has_lao:
                    spans.append(
                        f'<span class="word">'
                        f'<span class="orig">{orig_esc}</span>'
                        f'<span class="gloss">&nbsp;</span>'
                        f'<span class="lao">&nbsp;</span>'
                        f"</span>"
                    )
                else:
                    spans.append(
                        f'<span class="word">'
                        f'<span class="orig">{orig_esc}</span>'
                        f'<span class="gloss">&nbsp;</span>'
                        f"</span>"
                    )
            else:
                gloss = _html.escape(_format_ipa(ipa, ambiguous))
                css = "word ambig" if ambiguous else "word"
                if has_lao:
                    lao_esc = _html.escape(lao) if lao else "&nbsp;"
                    spans.append(
                        f'<span class="{css}">'
                        f'<span class="lao">{lao_esc}</span>'
                        f'<span class="gloss">{gloss}</span>'
                        f'<span class="orig">{orig_esc}</span>'
                        f"</span>"
                    )
                else:
                    spans.append(
                        f'<span class="{css}">'
                        f'<span class="orig">{orig_esc}</span>'
                        f'<span class="gloss">{gloss}</span>'
                        f"</span>"
                    )

        parts.append("<p>" + "".join(spans) + "</p>")

    parts += ["</body>", "</html>"]
    return "\n".join(parts).encode("utf-8")
